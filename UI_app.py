"""
Generative UI Builder
=====================
A Streamlit application that accepts multimodal inputs (text, CSV, images)
and uses an LLM to dynamically generate a structured UI layout in real time.

Author  : Assignment Submission
Model   : gpt-5.4-mini  (OpenAI)
"""

# =========================================================
# IMPORTS
# =========================================================

import os
import json
import base64
import textwrap
from io import BytesIO
from typing import Any, Dict, List, Optional

import ast
from dataclasses import dataclass

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from PIL import Image
from pydantic import BaseModel, ValidationError
from pypdf import PdfReader
from docx import Document as DocxDocument


@dataclass
class PythonFile:
    """Holds parsed metadata about an uploaded .py file."""
    source: str
    filename: str
    line_count: int
    functions: List[str]
    classes: List[str]
    imports: List[str]


@dataclass
class PdfFile:
    """Extracted content from an uploaded PDF."""
    filename: str
    page_count: int
    text: str          # full extracted text (may be truncated for prompt)
    metadata: Dict[str, Any]


@dataclass
class WordFile:
    """Extracted content from an uploaded .docx file."""
    filename: str
    paragraph_count: int
    text: str          # full plain-text content
    headings: List[str]
    tables_count: int

# =========================================================
# PAGE CONFIG  (must be the very first Streamlit call)
# =========================================================

st.set_page_config(
    page_title="Generative UI Builder",
    page_icon="✦",
    layout="wide",
    initial_sidebar_state="expanded",
)

# =========================================================
# CUSTOM CSS  – refined dark-editorial aesthetic
# =========================================================

st.markdown(
    """
    <style>
    /* ── Google Fonts ── */
    @import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500;600&display=swap');

    /* ── CSS Variables ── */
    :root {
        --bg-primary:   #0d0f14;
        --bg-surface:   #13161e;
        --bg-elevated:  #1a1e29;
        --border:       #252a38;
        --accent:       #6c8aff;
        --accent-warm:  #ff8c5a;
        --text-primary: #e8eaf0;
        --text-muted:   #7a8099;
        --text-faint:   #454c63;
        --success:      #5cba8e;
        --warning:      #f0a843;
        --error:        #e05c6b;
        --radius:       10px;
        --radius-sm:    6px;
        --font-display: 'DM Serif Display', Georgia, serif;
        --font-body:    'DM Sans', sans-serif;
        --font-mono:    'DM Mono', 'Courier New', monospace;
    }

    /* ── Global resets ── */
    html, body, [data-testid="stAppViewContainer"] {
        background-color: var(--bg-primary) !important;
        color: var(--text-primary) !important;
        font-family: var(--font-body) !important;
    }

    [data-testid="stSidebar"] {
        background-color: var(--bg-surface) !important;
        border-right: 1px solid var(--border) !important;
    }

    [data-testid="stHeader"] { background: transparent !important; }

    /* ── Hero header ── */
    .hero-header {
        padding: 2.5rem 0 1.5rem;
        border-bottom: 1px solid var(--border);
        margin-bottom: 2rem;
    }
    .hero-title {
        font-family: var(--font-display);
        font-size: 2.8rem;
        color: var(--text-primary);
        letter-spacing: -0.02em;
        line-height: 1.1;
        margin: 0;
    }
    .hero-title span { color: var(--accent); font-style: italic; }
    .hero-subtitle {
        font-family: var(--font-body);
        font-size: 0.95rem;
        color: var(--text-muted);
        margin-top: 0.5rem;
        font-weight: 300;
        letter-spacing: 0.01em;
    }

    /* ── Section label ── */
    .section-label {
        font-family: var(--font-mono);
        font-size: 0.7rem;
        letter-spacing: 0.15em;
        text-transform: uppercase;
        color: var(--text-faint);
        margin-bottom: 0.5rem;
    }

    /* ── Card ── */
    .gen-card {
        background: var(--bg-surface);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.5rem;
        margin-bottom: 1rem;
    }
    .gen-card:hover { border-color: var(--accent); transition: border-color .2s; }

    /* ── Metric card ── */
    .metric-card {
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.2rem 1.4rem;
        text-align: center;
    }
    .metric-card .metric-value {
        font-family: var(--font-display);
        font-size: 2.2rem;
        color: var(--accent);
        line-height: 1;
        display: block;
    }
    .metric-card .metric-label {
        font-size: 0.78rem;
        font-family: var(--font-mono);
        color: var(--text-muted);
        letter-spacing: 0.08em;
        text-transform: uppercase;
        margin-top: 0.35rem;
        display: block;
    }
    .metric-card .metric-delta {
        font-size: 0.8rem;
        margin-top: 0.25rem;
        display: block;
    }

    /* ── Badge ── */
    .badge {
        display: inline-block;
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-radius: 999px;
        padding: 0.15rem 0.65rem;
        font-size: 0.72rem;
        font-family: var(--font-mono);
        color: var(--text-muted);
        letter-spacing: 0.06em;
    }

    /* ── Status strip ── */
    .status-strip {
        display: flex;
        gap: 1rem;
        align-items: center;
        padding: 0.6rem 1rem;
        background: var(--bg-elevated);
        border-radius: var(--radius-sm);
        border: 1px solid var(--border);
        font-size: 0.82rem;
        color: var(--text-muted);
        margin-bottom: 1.5rem;
        font-family: var(--font-mono);
    }
    .status-strip .dot {
        width: 7px; height: 7px;
        border-radius: 50%;
        background: var(--success);
        box-shadow: 0 0 6px var(--success);
        flex-shrink: 0;
    }

    /* ── Divider ── */
    .section-divider {
        height: 1px;
        background: linear-gradient(90deg, var(--border) 0%, transparent 100%);
        margin: 2rem 0;
    }

    /* ── Streamlit widget overrides ── */
    .stTextArea textarea, .stTextInput input {
        background: var(--bg-elevated) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-sm) !important;
        color: var(--text-primary) !important;
        font-family: var(--font-body) !important;
        font-size: 0.9rem !important;
    }
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: var(--accent) !important;
        box-shadow: 0 0 0 2px rgba(108,138,255,.18) !important;
    }
    [data-testid="stFileUploader"] {
        background: var(--bg-elevated) !important;
        border: 1px dashed var(--border) !important;
        border-radius: var(--radius) !important;
    }
    .stButton > button {
        background: var(--accent) !important;
        color: #fff !important;
        border: none !important;
        border-radius: var(--radius-sm) !important;
        font-family: var(--font-body) !important;
        font-weight: 600 !important;
        letter-spacing: 0.02em !important;
        padding: 0.55rem 1.6rem !important;
        transition: opacity .2s, transform .15s !important;
    }
    .stButton > button:hover { opacity: 0.88 !important; transform: translateY(-1px) !important; }
    .stButton > button:active { transform: translateY(0) !important; }

    div[data-testid="stExpander"] {
        background: var(--bg-surface) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius) !important;
    }

    [data-testid="stMetric"] {
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1rem 1.2rem;
    }
    [data-testid="stMetricValue"] { color: var(--accent) !important; }

    /* ── Code block ── */
    .stCode, pre {
        background: var(--bg-elevated) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius) !important;
        font-family: var(--font-mono) !important;
        font-size: 0.82rem !important;
    }

    /* ── Alert boxes ── */
    .stAlert { border-radius: var(--radius-sm) !important; }

    /* ── Dataframe / table ── */
    [data-testid="stDataFrame"] { border-radius: var(--radius) !important; overflow: hidden; }

    /* ── Sidebar content ── */
    .sidebar-section {
        padding: 0.8rem 0;
        border-bottom: 1px solid var(--border);
        margin-bottom: 0.8rem;
    }
    .sidebar-heading {
        font-family: var(--font-mono);
        font-size: 0.68rem;
        letter-spacing: 0.14em;
        text-transform: uppercase;
        color: var(--text-faint);
        margin-bottom: 0.5rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# =========================================================
# ENV + CLIENT
# =========================================================

load_dotenv()

_api_key = os.getenv("OPENAI_API_KEY", "")

client = OpenAI(api_key=_api_key)

MODEL_NAME = "gpt-5.4-mini"

# =========================================================
# PYDANTIC SCHEMA
# =========================================================


class Component(BaseModel):
    type: str
    content: Dict[str, Any]


class UIResponse(BaseModel):
    title: str
    layout: str
    components: List[Component]


# =========================================================
# SYSTEM PROMPT
# =========================================================

SYSTEM_PROMPT = textwrap.dedent("""
    You are an expert UI architect and data visualisation specialist.

    Your sole task is to analyse user-provided content and return a
    STRICT JSON object that defines a frontend layout. No prose, no
    markdown code fences — raw JSON only.

    ─────────────────────────────────────────────
    REQUIRED OUTPUT STRUCTURE
    ─────────────────────────────────────────────
    {
      "title": "<descriptive dashboard title>",
      "layout": "wide",
      "components": [
        {
          "type": "<see types below>",
          "content": { ... }
        }
      ]
    }

    ─────────────────────────────────────────────
    COMPONENT TYPES & CONTENT SCHEMAS
    ─────────────────────────────────────────────
    text
      { "text": "<markdown string>" }

    metric
      { "label": "...", "value": "...", "delta": "..." }
      (delta is optional; prefix '+' / '-' to show direction)

    table
      { "data": { "col1": [v1,v2,...], "col2": [v1,v2,...] },
        "caption": "optional caption" }

    bar_chart
      { "data": { "<x-axis label>": [v1,v2,...], "<y-col>": [v1,v2,...] },
        "x": "<column name for x-axis>",
        "y": ["<y-col1>", "<y-col2>"],   ← optional, list of y columns
        "title": "...",
        "orientation": "v"  ← "v" (default) or "h" for horizontal }

    line_chart
      { "data": { "<x-axis label>": [v1,v2,...], "<y-col>": [v1,v2,...] },
        "x": "<column name for x-axis>",
        "y": ["<y-col>"],
        "title": "..." }

    area_chart
      { "data": { "<x-axis label>": [v1,v2,...], "<y-col>": [v1,v2,...] },
        "x": "<column name for x-axis>",
        "y": ["<y-col>"],
        "title": "..." }

    pie_chart
      { "data": { "Label A": 42, "Label B": 58, ... },
        "title": "...",
        "hole": 0.0 }    ← 0.0 = full pie; 0.4 = donut chart

    scatter_chart
      { "data": { "<x-col>": [v1,v2,...], "<y-col>": [v1,v2,...],
                  "<color-col>": ["cat1","cat2",...] },  ← color optional
        "x": "<x-col>",
        "y": "<y-col>",
        "color": "<color-col>",   ← optional
        "title": "..." }

    histogram
      { "data": { "<col>": [v1,v2,...] },
        "x": "<col>",
        "nbins": 20,    ← optional
        "title": "..." }

    image
      { "url": "<public image URL>", "caption": "..." }

    ─────────────────────────────────────────────
    CHART SELECTION RULES  ← CRITICAL: always follow these
    ─────────────────────────────────────────────
    The data summary you receive contains a "chart_suitability_hints" array.
    Read it carefully and apply these rules:

    1. time_series_detected      → USE line_chart or area_chart (area for cumulative/volume data)
       Set x to the datetime column. Never use bar_chart for time series.

    2. multiple_numeric_cols     → USE scatter_chart to reveal correlations between two numeric columns.
       Also add a line_chart showing each numeric column over index/time if applicable.

    3. low_cardinality_categorical (≤ 12 unique values) paired with a numeric:
       → If ≤ 6 categories: USE pie_chart (hole: 0.4 for donut looks better).
       → If 7–12 categories: USE bar_chart (horizontal orientation if labels are long).

    4. high_cardinality_categorical (> 12 unique values):
       → USE bar_chart (top-N bar chart, limit to 15 bars) + table for full data.

    5. continuous_distribution detected → USE histogram for that column.

    6. Single numeric column, no categories → USE histogram.

    7. Multiple numeric columns with no time axis → USE bar_chart (grouped) + scatter_chart.

    8. NEVER produce only a table when chart data is available.
    9. ALWAYS include at least one chart component for CSV/tabular data.
    10. For mixed data: combine the best chart + a table + metrics.

    ─────────────────────────────────────────────
    GENERAL DECISION RULES
    ─────────────────────────────────────────────
    • Pure text prompt (no file) → Build a complete dashboard from the prompt topic.
      Produce: 1 text overview, 3-5 metric cards with plausible values,
      1-2 charts with realistic sample data, and 1 table.
    • CSV / tabular data → Apply chart selection rules above. Use only numbers
      from the supplied data. Produce: 2-4 metrics, 1 table, 1-2 charts, text summary.
    • Image input      → text description component, insights text block, metric cards.
    • Plain text file  → structured text summary, metrics if numbers present.
    • Python file      → metric cards for code stats, text overview, function/class table.
    • PDF / Word file  → text summary, metric cards, table of headings/sections.
    • Always start with an insightful text component as the first item.
    • Keep titles concise and meaningful.
    • Return ONLY the JSON object. Nothing else.
""")

# =========================================================
# HELPERS – FILE READING
# =========================================================


def read_uploaded_file(uploaded_file) -> Optional[Any]:
    """Return parsed content (DataFrame, PIL Image, str, or None)."""
    if uploaded_file is None:
        return None

    file_type = uploaded_file.type

    if file_type == "text/plain":
        return uploaded_file.read().decode("utf-8")

    if file_type == "text/csv":
        return pd.read_csv(uploaded_file)

    if file_type.startswith("image/"):
        return Image.open(uploaded_file)

    # ── Python source file ──
    if file_type in ("text/x-python", "application/x-python", "text/x-python-script") or \
            (uploaded_file.name or "").endswith(".py"):
        return parse_python_file(uploaded_file)

    # ── PDF ──
    if file_type == "application/pdf" or (uploaded_file.name or "").endswith(".pdf"):
        return parse_pdf_file(uploaded_file)

    # ── Word (.docx) ──
    if file_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or (uploaded_file.name or "").endswith(".docx"):
        return parse_word_file(uploaded_file)

    return None


def parse_python_file(uploaded_file) -> "PythonFile":
    """Read a .py file and extract structural metadata via AST."""
    source = uploaded_file.read().decode("utf-8")
    filename = uploaded_file.name or "script.py"
    line_count = source.count("\n") + 1

    functions: List[str] = []
    classes: List[str] = []
    imports: List[str] = []

    try:
        tree = ast.parse(source)
        for node in ast.walk(tree):
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                functions.append(node.name)
            elif isinstance(node, ast.ClassDef):
                classes.append(node.name)
            elif isinstance(node, ast.Import):
                for alias in node.names:
                    imports.append(alias.name)
            elif isinstance(node, ast.ImportFrom):
                module = node.module or ""
                for alias in node.names:
                    imports.append(f"{module}.{alias.name}" if module else alias.name)
    except SyntaxError:
        pass  # Still return what we have; prompt will note parse failure

    return PythonFile(
        source=source,
        filename=filename,
        line_count=line_count,
        functions=functions,
        classes=classes,
        imports=list(dict.fromkeys(imports)),  # deduplicate, preserve order
    )


def parse_pdf_file(uploaded_file) -> "PdfFile":
    """Extract text and metadata from a PDF using pypdf."""
    filename = uploaded_file.name or "document.pdf"
    raw_bytes = uploaded_file.read()
    reader = PdfReader(BytesIO(raw_bytes))

    pages_text: List[str] = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            pages_text.append("")

    full_text = "\n\n".join(pages_text)

    # pypdf metadata values can be IndirectObject; cast to str safely
    raw_meta = reader.metadata or {}
    metadata: Dict[str, Any] = {}
    for k, v in raw_meta.items():
        try:
            metadata[str(k)] = str(v)
        except Exception:
            pass

    return PdfFile(
        filename=filename,
        page_count=len(reader.pages),
        text=full_text,
        metadata=metadata,
    )


def parse_word_file(uploaded_file) -> "WordFile":
    """Extract text, headings, and table count from a .docx file."""
    filename = uploaded_file.name or "document.docx"
    raw_bytes = uploaded_file.read()
    doc = DocxDocument(BytesIO(raw_bytes))

    paragraphs: List[str] = []
    headings: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        if para.style.name.startswith("Heading"):
            headings.append(text)

    full_text = "\n\n".join(paragraphs)

    return WordFile(
        filename=filename,
        paragraph_count=len(paragraphs),
        text=full_text,
        headings=headings,
        tables_count=len(doc.tables),
    )


def dataframe_to_prompt_json(df: pd.DataFrame) -> str:
    """Serialise a DataFrame into a compact summary for the LLM prompt,
    including chart-suitability hints so the model picks the right chart type."""
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    categorical_cols = df.select_dtypes(include=["object", "category"]).columns.tolist()
    datetime_cols = df.select_dtypes(include=["datetime", "datetimetz"]).columns.tolist()

    # Detect likely datetime columns stored as strings (e.g. "2024-01", "Jan 2024")
    for col in categorical_cols:
        try:
            pd.to_datetime(df[col], infer_datetime_format=True)
            datetime_cols.append(col)
        except Exception:
            pass

    # Cardinality of categorical columns
    cardinality = {c: int(df[c].nunique()) for c in categorical_cols}

    # Chart suitability hints
    hints = []
    if datetime_cols and numeric_cols:
        hints.append("time_series_detected: prefer line_chart or area_chart with x=" + datetime_cols[0])
    if len(numeric_cols) >= 2:
        hints.append("multiple_numeric_cols: scatter_chart suitable for correlation")
    if categorical_cols and numeric_cols:
        low_card = [c for c in categorical_cols if cardinality.get(c, 99) <= 12]
        if low_card:
            hints.append(f"low_cardinality_categorical ({low_card[0]}, {cardinality[low_card[0]]} values): pie_chart or bar_chart suitable")
        else:
            hints.append("high_cardinality_categorical: prefer bar_chart or table")
    if all(df[c].nunique() / max(len(df), 1) > 0.8 for c in numeric_cols[:1]):
        hints.append("continuous_distribution: histogram suitable")

    numeric_summary = df.describe(include="number").round(2).to_dict()
    summary = {
        "columns": list(df.columns),
        "shape": {"rows": int(df.shape[0]), "columns": int(df.shape[1])},
        "dtypes": {c: str(t) for c, t in df.dtypes.items()},
        "numeric_columns": numeric_cols,
        "categorical_columns": categorical_cols,
        "datetime_columns": datetime_cols,
        "categorical_cardinality": cardinality,
        "chart_suitability_hints": hints,
        "sample_rows": df.head(8).to_dict(orient="records"),
        "numeric_statistics": numeric_summary,
    }
    return json.dumps(summary, default=str)


def image_to_base64(image: Image.Image, max_size: int = 512) -> str:
    """Resize and base64-encode a PIL image for multimodal API calls."""
    image.thumbnail((max_size, max_size), Image.LANCZOS)
    buf = BytesIO()
    fmt = image.format or "PNG"
    if fmt not in ("PNG", "JPEG", "WEBP", "GIF"):
        fmt = "PNG"
    image.save(buf, format=fmt)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def build_messages(
    text_input: str,
    uploaded_content,
    temperature: float,
) -> List[Dict]:
    """Construct the messages array to send to the API."""

    messages: List[Dict] = [{"role": "system", "content": SYSTEM_PROMPT}]

    user_parts: List[Dict] = []

    # ── Text prompt ──
    if text_input.strip():
        user_parts.append({"type": "text", "text": f"User prompt:\n{text_input.strip()}"})

    # ── CSV ──
    if isinstance(uploaded_content, pd.DataFrame):
        user_parts.append(
            {
                "type": "text",
                "text": (
                    "The user uploaded a CSV file. "
                    "Here is a detailed statistical summary:\n"
                    + dataframe_to_prompt_json(uploaded_content)
                ),
            }
        )

    # ── Image (multimodal) ──
    elif isinstance(uploaded_content, Image.Image):
        mime = "image/png"
        if uploaded_content.format == "JPEG":
            mime = "image/jpeg"
        elif uploaded_content.format == "WEBP":
            mime = "image/webp"

        b64 = image_to_base64(uploaded_content)
        user_parts.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "low"},
            }
        )
        user_parts.append(
            {"type": "text", "text": "Analyse this image and build an appropriate UI."}
        )

    # ── Plain text file ──
    elif isinstance(uploaded_content, str):
        excerpt = uploaded_content[:4000]
        user_parts.append(
            {
                "type": "text",
                "text": f"The user uploaded a text file. Here is an excerpt:\n\n{excerpt}",
            }
        )

    # ── Python source file ──
    elif isinstance(uploaded_content, PythonFile):
        pf = uploaded_content
        code_excerpt = pf.source[:5000]
        truncation_note = (
            f"\n\n[Source truncated — showing first 5 000 of {len(pf.source)} chars]"
            if len(pf.source) > 5000 else ""
        )
        summary = (
            f"Filename   : {pf.filename}\n"
            f"Lines      : {pf.line_count}\n"
            f"Classes    : {', '.join(pf.classes) or 'none'}\n"
            f"Functions  : {', '.join(pf.functions) or 'none'}\n"
            f"Imports    : {', '.join(pf.imports[:20]) or 'none'}"
            + (" …" if len(pf.imports) > 20 else "")
        )
        user_parts.append(
            {
                "type": "text",
                "text": (
                    "The user uploaded a Python source file.\n\n"
                    f"── Structural Summary ──\n{summary}\n\n"
                    f"── Source Code ──\n```python\n{code_excerpt}{truncation_note}\n```\n\n"
                    "Analyse this code and build a UI that documents and visualises it: "
                    "include a text overview, metric cards for key stats (lines, functions, "
                    "classes, imports), a table of functions/classes, and any other relevant components."
                ),
            }
        )

    # ── PDF file ──
    elif isinstance(uploaded_content, PdfFile):
        pdf = uploaded_content
        excerpt = pdf.text[:6000]
        truncation_note = (
            f"\n\n[Text truncated — showing first 6 000 of {len(pdf.text)} chars]"
            if len(pdf.text) > 6000 else ""
        )
        meta_str = ", ".join(f"{k}: {v}" for k, v in pdf.metadata.items()) or "none"
        user_parts.append(
            {
                "type": "text",
                "text": (
                    f"The user uploaded a PDF file: {pdf.filename}\n"
                    f"Pages: {pdf.page_count}  |  Metadata: {meta_str}\n\n"
                    f"── Extracted Text ──\n{excerpt}{truncation_note}\n\n"
                    "Analyse this document and build an appropriate UI: include a text summary, "
                    "key metric cards (page count, word count, etc.), and any relevant charts or "
                    "tables if the content contains structured data or statistics."
                ),
            }
        )

    # ── Word (.docx) file ──
    elif isinstance(uploaded_content, WordFile):
        wf = uploaded_content
        excerpt = wf.text[:6000]
        truncation_note = (
            f"\n\n[Text truncated — showing first 6 000 of {len(wf.text)} chars]"
            if len(wf.text) > 6000 else ""
        )
        headings_str = "\n".join(f"  • {h}" for h in wf.headings[:30]) or "  (none detected)"
        user_parts.append(
            {
                "type": "text",
                "text": (
                    f"The user uploaded a Word document: {wf.filename}\n"
                    f"Paragraphs: {wf.paragraph_count}  |  Tables: {wf.tables_count}\n\n"
                    f"── Document Headings ──\n{headings_str}\n\n"
                    f"── Document Text ──\n{excerpt}{truncation_note}\n\n"
                    "Analyse this document and build an appropriate UI: include a text summary, "
                    "metric cards (paragraph count, table count, section count), and a table of "
                    "headings/sections. Add charts if the content contains data or statistics."
                ),
            }
        )

    # ── Pure text prompt (no file) — already added above via text_input.strip()
    # The fallback below only fires if Generate UI is somehow called with nothing at all.
    if not user_parts:
        user_parts.append(
            {"type": "text", "text": "Generate a sample KPI dashboard as a demonstration."}
        )

    messages.append({"role": "user", "content": user_parts})
    return messages


# =========================================================
# PYTHON-SPECIFIC LLM HELPERS
# =========================================================


def generate_completed_code(pf: "PythonFile", temperature: float = 0.2) -> Optional[str]:
    """
    Send the Python source to the model and ask it to fill every gap:
      - functions / methods whose body is only `pass` or `...`
      - missing docstrings on functions, classes and the module
      - inline TODO / FIXME / HACK comments
      - type annotations that are absent
    Returns the completed source as a plain string, or None on error.
    """
    system = textwrap.dedent("""
        You are an expert Python engineer and code-completion assistant.

        The user will provide a Python source file that may contain gaps such as:
          • Functions or methods whose body is only `pass`, `...`, or a bare `raise NotImplementedError`
          • Missing docstrings (module, class, function level)
          • TODO / FIXME / HACK / XXX comments that indicate unfinished logic
          • Missing type annotations on function signatures

        Your task:
        1. Fill in every gap with correct, idiomatic Python that fits the surrounding context.
        2. Add or improve docstrings (Google style) where they are absent or minimal.
        3. Add type annotations to all function signatures that lack them.
        4. Resolve every TODO/FIXME/HACK comment with real implementation.
        5. Do NOT change any logic that is already implemented.
        6. Do NOT add new imports unless absolutely required by your additions.
        7. Return ONLY the completed Python source code — no prose, no markdown fences.
    """)

    user_msg = (
        f"File: {pf.filename}\n\n"
        f"```python\n{pf.source}\n```\n\n"
        "Return the fully completed Python source file."
    )

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user_msg},
            ],
            temperature=temperature,
        )
        raw = response.choices[0].message.content or ""
        # Strip accidental markdown fences
        raw = raw.strip()
        if raw.startswith("```"):
            lines = raw.splitlines()
            end = next(
                (i for i, l in enumerate(lines[1:], 1) if l.strip() == "```"),
                len(lines),
            )
            raw = "\n".join(lines[1:end])
        return raw
    except Exception as exc:
        st.error(f"**Code-completion API Error:** {exc}")
        return None


def generate_markdown_doc(pf: "PythonFile", temperature: float = 0.2) -> Optional[str]:
    """
    Ask the model to produce a full Markdown documentation file for the
    Python source — including overview, installation, usage, API reference
    (every class and function), and examples.
    Returns the markdown string, or None on error.
    """
    system = textwrap.dedent("""
        You are a technical writer who specialises in Python library documentation.

        Given a Python source file, produce a comprehensive Markdown documentation
        file with the following sections (use ## headings):

        1. Title & short description (use the module docstring if present)
        2. ## Overview  – what the module does and its architecture
        3. ## Requirements  – dependencies inferred from imports
        4. ## Installation  – generic pip install instructions
        5. ## Usage  – concise code examples showing typical use
        6. ## API Reference  – one ### subsection per class and per top-level function:
             - Signature with type hints
             - Purpose (one sentence)
             - Parameters table (Name | Type | Description)
             - Returns row
             - Raises (if applicable)
             - Short example
        7. ## Notes  – design decisions, caveats, known limitations

        Rules:
        - Use fenced code blocks with `python` syntax highlighting for all examples.
        - Be specific — derive everything from the actual source, do not invent behaviour.
        - Return ONLY the Markdown content, no extra commentary.
    """)

    user_msg = (
        f"File: {pf.filename}\n\n"
        f"```python\n{pf.source}\n```\n\n"
        "Produce the full Markdown documentation."
    )

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user_msg},
            ],
            temperature=temperature,
        )
        return response.choices[0].message.content or ""
    except Exception as exc:
        st.error(f"**Markdown-doc API Error:** {exc}")
        return None


# =========================================================
# OPENAI STREAMING CALL
# =========================================================


def generate_ui_json(
    text_input: str,
    uploaded_content,
    temperature: float = 0.3,
) -> Optional[str]:
    """
    Call the model with streaming enabled.
    Displays a live JSON preview in the sidebar and returns the full response.
    """
    messages = build_messages(text_input, uploaded_content, temperature)

    try:
        stream = client.chat.completions.create(
            model=MODEL_NAME,
            messages=messages,
            temperature=temperature,
            stream=True,
        )
    except Exception as exc:
        st.error(f"**API Error:** {exc}")
        return None

    full_response = ""
    stream_box = st.empty()

    for chunk in stream:
        delta = chunk.choices[0].delta
        if delta and delta.content:
            full_response += delta.content
            # Live preview – truncate to last 1 200 chars to stay tidy
            preview = full_response if len(full_response) < 1200 else "…" + full_response[-1200:]
            stream_box.code(preview, language="json")

    stream_box.empty()  # Clear streaming preview once done
    return full_response


# =========================================================
# JSON PARSING + VALIDATION
# =========================================================


def parse_ui_response(raw: str) -> Optional[UIResponse]:
    """Parse and validate the LLM output with helpful error messages."""
    # Strip possible markdown fences the model might accidentally include
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        lines = cleaned.splitlines()
        cleaned = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        st.error(
            f"**JSON Parse Error** – The model returned malformed JSON.\n\n"
            f"`{exc}`\n\n"
            "Try regenerating or simplifying your prompt."
        )
        return None

    try:
        return UIResponse(**parsed)
    except ValidationError as exc:
        st.error(f"**Schema Validation Error**\n\n```\n{exc}\n```")
        return None


# =========================================================
# DYNAMIC UI RENDERER
# =========================================================


def render_component(component: Component, idx: int) -> None:
    """Render a single UI component based on its type using Plotly for charts."""

    ctype = component.type.lower().strip()
    content = component.content

    # Shared Plotly theme
    PLOTLY_THEME = "plotly_dark"
    CHART_BG = "rgba(26,30,41,0)"       # transparent – matches --bg-surface
    PAPER_BG = "rgba(19,22,30,0)"
    FONT_COLOR = "#e8eaf0"
    GRID_COLOR = "#252a38"
    ACCENT = "#6c8aff"

    def _apply_base_layout(fig):
        fig.update_layout(
            template=PLOTLY_THEME,
            plot_bgcolor=CHART_BG,
            paper_bgcolor=PAPER_BG,
            font=dict(color=FONT_COLOR, family="DM Sans, sans-serif", size=12),
            xaxis=dict(gridcolor=GRID_COLOR, zerolinecolor=GRID_COLOR),
            yaxis=dict(gridcolor=GRID_COLOR, zerolinecolor=GRID_COLOR),
            margin=dict(l=10, r=10, t=40, b=10),
            legend=dict(bgcolor="rgba(0,0,0,0)", bordercolor=GRID_COLOR),
        )
        return fig

    # ── Text / Markdown ──────────────────────────────────
    if ctype == "text":
        st.markdown(content.get("text", ""))

    # ── Metric ───────────────────────────────────────────
    elif ctype == "metric":
        delta = content.get("delta")
        st.metric(
            label=content.get("label", "Metric"),
            value=str(content.get("value", "–")),
            delta=str(delta) if delta is not None else None,
        )

    # ── Table ────────────────────────────────────────────
    elif ctype == "table":
        data = content.get("data", {})
        caption = content.get("caption", "")
        try:
            df = pd.DataFrame(data)
            if caption:
                st.caption(caption)
            st.dataframe(df, use_container_width=True)
        except Exception:
            st.warning("⚠️ Could not render table – malformed data.")

    # ── Bar Chart (Plotly) ───────────────────────────────
    elif ctype == "bar_chart":
        data = content.get("data", {})
        title = content.get("title", "")
        x_col = content.get("x")
        y_cols = content.get("y")
        orientation = content.get("orientation", "v")
        try:
            df = pd.DataFrame(data)
            if x_col and x_col not in df.columns:
                x_col = None
            if not x_col:
                x_col = df.columns[0]
            if not y_cols:
                y_cols = [c for c in df.columns if c != x_col]
            elif isinstance(y_cols, str):
                y_cols = [y_cols]
            y_cols = [c for c in y_cols if c in df.columns]
            if not y_cols:
                y_cols = [c for c in df.columns if c != x_col]

            if orientation == "h":
                fig = px.bar(df, x=y_cols[0] if y_cols else None, y=x_col,
                             barmode="group", title=title, orientation="h",
                             color_discrete_sequence=px.colors.qualitative.Bold)
            else:
                fig = px.bar(df, x=x_col, y=y_cols, barmode="group", title=title,
                             color_discrete_sequence=px.colors.qualitative.Bold)

            _apply_base_layout(fig)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render bar chart: {e}")

    # ── Line Chart (Plotly) ──────────────────────────────
    elif ctype == "line_chart":
        data = content.get("data", {})
        title = content.get("title", "")
        x_col = content.get("x")
        y_cols = content.get("y")
        try:
            df = pd.DataFrame(data)
            if x_col and x_col not in df.columns:
                x_col = None
            if not x_col:
                x_col = df.columns[0]
            if not y_cols:
                y_cols = [c for c in df.columns if c != x_col]
            elif isinstance(y_cols, str):
                y_cols = [y_cols]
            y_cols = [c for c in y_cols if c in df.columns]

            fig = px.line(df, x=x_col, y=y_cols, title=title, markers=True,
                          color_discrete_sequence=px.colors.qualitative.Bold)
            _apply_base_layout(fig)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render line chart: {e}")

    # ── Area Chart (Plotly) ──────────────────────────────
    elif ctype == "area_chart":
        data = content.get("data", {})
        title = content.get("title", "")
        x_col = content.get("x")
        y_cols = content.get("y")
        try:
            df = pd.DataFrame(data)
            if x_col and x_col not in df.columns:
                x_col = None
            if not x_col:
                x_col = df.columns[0]
            if not y_cols:
                y_cols = [c for c in df.columns if c != x_col]
            elif isinstance(y_cols, str):
                y_cols = [y_cols]
            y_cols = [c for c in y_cols if c in df.columns]

            fig = px.area(df, x=x_col, y=y_cols, title=title,
                          color_discrete_sequence=px.colors.qualitative.Bold)
            _apply_base_layout(fig)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render area chart: {e}")

    # ── Pie / Donut Chart (Plotly – real pie!) ───────────
    elif ctype == "pie_chart":
        data = content.get("data", {})
        title = content.get("title", "")
        hole = float(content.get("hole", 0.0))
        try:
            labels = list(data.keys())
            values = list(data.values())
            fig = px.pie(names=labels, values=values, title=title, hole=hole,
                         color_discrete_sequence=px.colors.qualitative.Bold)
            fig.update_traces(textposition="inside", textinfo="percent+label",
                              textfont_color="white")
            _apply_base_layout(fig)
            fig.update_layout(showlegend=True)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render pie chart: {e}")

    # ── Scatter Chart (Plotly) ───────────────────────────
    elif ctype == "scatter_chart":
        data = content.get("data", {})
        title = content.get("title", "")
        x_col = content.get("x")
        y_col = content.get("y")
        color_col = content.get("color")
        try:
            df = pd.DataFrame(data)
            if not x_col or x_col not in df.columns:
                x_col = df.columns[0]
            if not y_col or y_col not in df.columns:
                numeric = df.select_dtypes(include="number").columns
                y_col = numeric[1] if len(numeric) > 1 else df.columns[1]
            color_arg = color_col if (color_col and color_col in df.columns) else None

            fig = px.scatter(df, x=x_col, y=y_col, color=color_arg, title=title,
                             trendline="ols" if color_arg is None else None,
                             color_discrete_sequence=px.colors.qualitative.Bold)
            _apply_base_layout(fig)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render scatter chart: {e}")

    # ── Histogram (Plotly) ───────────────────────────────
    elif ctype == "histogram":
        data = content.get("data", {})
        title = content.get("title", "")
        x_col = content.get("x")
        nbins = content.get("nbins", 20)
        try:
            df = pd.DataFrame(data)
            if not x_col or x_col not in df.columns:
                x_col = df.columns[0]
            fig = px.histogram(df, x=x_col, nbins=int(nbins), title=title,
                               color_discrete_sequence=[ACCENT])
            fig.update_traces(marker_line_color=GRID_COLOR, marker_line_width=0.5)
            _apply_base_layout(fig)
            st.plotly_chart(fig, use_container_width=True)
        except Exception as e:
            st.warning(f"⚠️ Could not render histogram: {e}")

    # ── Image ────────────────────────────────────────────
    elif ctype == "image":
        url = content.get("url", "")
        caption = content.get("caption", "")
        if url:
            st.image(url, caption=caption, use_column_width=True)
        else:
            st.info("No image URL provided.")

    # ── Unknown ──────────────────────────────────────────
    else:
        st.markdown(
            f'<div class="gen-card">'
            f'<span class="badge">unknown type: {ctype}</span>'
            f"<pre>{json.dumps(content, indent=2)}</pre>"
            f"</div>",
            unsafe_allow_html=True,
        )


# =========================================================
# SIDEBAR
# =========================================================


def render_sidebar() -> Dict[str, Any]:
    """Render sidebar controls and return settings dict."""

    with st.sidebar:
        st.markdown(
            '<div class="sidebar-heading">Model Settings</div>', unsafe_allow_html=True
        )
        temperature = st.slider(
            "Temperature", min_value=0.0, max_value=1.0, value=0.3, step=0.05,
            help="Lower = more deterministic JSON; higher = more creative layouts."
        )
        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

        st.markdown('<div class="sidebar-heading">About</div>', unsafe_allow_html=True)
        st.markdown(
            '<p style="font-size:0.82rem;color:var(--text-muted);line-height:1.55;">'
            "Upload a CSV, image, or plain-text file — or just type a prompt — "
            "and let the model design a custom dashboard in real time."
            "</p>",
            unsafe_allow_html=True,
        )

        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

        st.markdown(
            '<div class="sidebar-heading">Supported Components</div>',
            unsafe_allow_html=True,
        )
        for comp in ["text", "metric", "table", "bar_chart", "line_chart", "area_chart", "pie_chart", "scatter_chart", "histogram", "image"]:
            st.markdown(
                f'<span class="badge">{comp}</span>&nbsp;',
                unsafe_allow_html=True,
            )

    return {"temperature": temperature}


# =========================================================
# MAIN
# =========================================================


def main() -> None:
    settings = render_sidebar()

    # ── Hero ──────────────────────────────────────────────
    st.markdown(
        """
        <div class="hero-header">
            <h1 class="hero-title">Generative <span>UI</span> Builder</h1>
            <p class="hero-subtitle">
                Multimodal input · Structured JSON · Smart chart selection · Streaming
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ── Status strip ──────────────────────────────────────
    model_status = "connected" if _api_key else "⚠ API key not found"
    st.markdown(
        f"""
        <div class="status-strip">
            <span class="dot"></span>
            <span>Model: <strong>{MODEL_NAME}</strong></span>
            <span style="color:var(--text-faint)">|</span>
            <span>Key: <strong>{model_status}</strong></span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ── Input columns ─────────────────────────────────────
    col_left, col_right = st.columns([1.1, 0.9], gap="large")

    with col_left:
        st.markdown('<div class="section-label">Prompt</div>', unsafe_allow_html=True)
        text_input = st.text_area(
            label="",
            placeholder=(
                "e.g. 'Build a sales performance dashboard' "
                "or describe what you want the interface to show…"
            ),
            height=140,
            label_visibility="collapsed",
        )

    with col_right:
        st.markdown('<div class="section-label">File Upload (optional)</div>', unsafe_allow_html=True)
        uploaded_file = st.file_uploader(
            label="",
            type=["txt", "csv", "png", "jpg", "jpeg", "webp", "py", "pdf", "docx"],
            label_visibility="collapsed",
        )

    # ── Preview ───────────────────────────────────────────
    uploaded_content = read_uploaded_file(uploaded_file)

    if uploaded_content is not None:
        with st.expander("📂 File Preview", expanded=False):
            if isinstance(uploaded_content, pd.DataFrame):
                st.caption(
                    f"**{uploaded_content.shape[0]:,} rows × {uploaded_content.shape[1]} columns**"
                )
                st.dataframe(uploaded_content.head(10), use_container_width=True)

            elif isinstance(uploaded_content, Image.Image):
                st.image(uploaded_content, use_column_width=True)

            elif isinstance(uploaded_content, str):
                st.code(uploaded_content[:2000], language="text")

            elif isinstance(uploaded_content, PythonFile):
                pf = uploaded_content
                st.caption(
                    f"**{pf.filename}** · {pf.line_count} lines · "
                    f"{len(pf.functions)} functions · {len(pf.classes)} classes · "
                    f"{len(pf.imports)} imports"
                )
                st.code(pf.source[:3000], language="python")

            elif isinstance(uploaded_content, PdfFile):
                pdf = uploaded_content
                st.caption(
                    f"**{pdf.filename}** · {pdf.page_count} pages · "
                    f"{len(pdf.text.split()):,} words"
                )
                if pdf.metadata:
                    st.json({k: v for k, v in list(pdf.metadata.items())[:6]})
                st.text(pdf.text[:3000] + ("…" if len(pdf.text) > 3000 else ""))

            elif isinstance(uploaded_content, WordFile):
                wf = uploaded_content
                st.caption(
                    f"**{wf.filename}** · {wf.paragraph_count} paragraphs · "
                    f"{wf.tables_count} tables · {len(wf.headings)} headings"
                )
                if wf.headings:
                    st.markdown("**Headings detected:**")
                    for h in wf.headings[:15]:
                        st.markdown(f"- {h}")
                st.text(wf.text[:3000] + ("…" if len(wf.text) > 3000 else ""))

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ── Python auto-processing (runs as soon as a .py file is uploaded) ───────
    if isinstance(uploaded_content, PythonFile):
        pf = uploaded_content
        st.markdown(
            '<div class="section-label">Python File — Auto Processing</div>',
            unsafe_allow_html=True,
        )

        col_run1, col_run2, _ = st.columns([0.22, 0.22, 0.56])
        with col_run1:
            run_completion = st.button("⚙ Fill Code Gaps", use_container_width=True)
        with col_run2:
            run_docs = st.button("📄 Generate Docs", use_container_width=True)

        # ── Fill gaps ──────────────────────────────────────────────────────────
        if run_completion:
            if not _api_key:
                st.error("Set `OPENAI_API_KEY` before running.")
            else:
                with st.spinner("Filling gaps in the code…"):
                    completed = generate_completed_code(pf, temperature=settings["temperature"])

                if completed:
                    st.success("✔ Code gaps filled successfully.")
                    tab_diff, tab_full = st.tabs(["Completed Code", "Download"])

                    with tab_diff:
                        st.code(completed, language="python")

                    with tab_full:
                        base_name = pf.filename.replace(".py", "")
                        st.download_button(
                            label="⬇ Download completed .py",
                            data=completed.encode("utf-8"),
                            file_name=f"{base_name}_completed.py",
                            mime="text/x-python",
                            use_container_width=True,
                        )

        # ── Generate markdown docs ─────────────────────────────────────────────
        if run_docs:
            if not _api_key:
                st.error("Set `OPENAI_API_KEY` before running.")
            else:
                with st.spinner("Writing documentation…"):
                    md_doc = generate_markdown_doc(pf, temperature=settings["temperature"])

                if md_doc:
                    st.success("✔ Documentation generated successfully.")
                    tab_preview, tab_dl = st.tabs(["Preview", "Download"])

                    with tab_preview:
                        st.markdown(md_doc)

                    with tab_dl:
                        base_name = pf.filename.replace(".py", "")
                        st.download_button(
                            label="⬇ Download .md file",
                            data=md_doc.encode("utf-8"),
                            file_name=f"{base_name}_docs.md",
                            mime="text/markdown",
                            use_container_width=True,
                        )

        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ── Generate button ────────────────────────────────────
    btn_col, _ = st.columns([0.2, 0.8])
    with btn_col:
        generate = st.button("✦ Generate UI", use_container_width=True)

    # ── Generation + render ────────────────────────────────
    if generate:
        if not _api_key:
            st.error("Set `OPENAI_API_KEY` in your environment or `.env` file before generating.")
            return

        if not text_input.strip() and uploaded_content is None:
            st.warning("Please enter a prompt or upload a file to continue.")
            return

        with st.spinner("Calling model…"):
            raw_response = generate_ui_json(
                text_input,
                uploaded_content,
                temperature=settings["temperature"],
            )

        if raw_response is None:
            return

        parsed_ui = parse_ui_response(raw_response)

        if parsed_ui is None:
            return

        # ── Rendered output ────────────────────────────────
        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
        st.markdown(
            f'<h2 style="font-family:var(--font-display);font-size:1.9rem;'
            f'letter-spacing:-0.02em;margin-bottom:1.2rem;">'
            f"{parsed_ui.title}</h2>",
            unsafe_allow_html=True,
        )

        # Group metrics together for a clean grid
        metrics = [c for c in parsed_ui.components if c.type.lower() == "metric"]
        non_metrics = [c for c in parsed_ui.components if c.type.lower() != "metric"]

        if metrics:
            m_cols = st.columns(min(len(metrics), 4))
            for i, mc in enumerate(metrics):
                with m_cols[i % len(m_cols)]:
                    render_component(mc, i)

        for idx, component in enumerate(non_metrics):
            render_component(component, idx)

        # ── Raw JSON expander ──────────────────────────────
        with st.expander("{ } View Raw JSON Schema", expanded=False):
            st.json(parsed_ui.model_dump())


# =========================================================
# ENTRY POINT
# =========================================================

if __name__ == "__main__":
    main()